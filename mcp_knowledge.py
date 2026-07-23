"""
📚 知识库 MCP 服务 - 自动读取 knowledge_base 文件夹中的文件
支持：.txt / .xlsx / .xls / .docx
用户直接把文件丢进 knowledge_base 文件夹，Bot 自动识别读取
"""
from mcp.server.fastmcp import FastMCP
import os
import glob

mcp = FastMCP("本地知识库服务")

# ── 知识库路径（与 mcp 文件同目录下的 knowledge_base 文件夹） ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
KB_DIR = os.path.join(BASE_DIR, "knowledge_base")


def ensure_kb_dir():
    """确保知识库文件夹存在"""
    if not os.path.exists(KB_DIR):
        os.makedirs(KB_DIR)


def get_supported_files() -> list[str]:
    """扫描知识库文件夹，返回所有支持的文件路径"""
    ensure_kb_dir()
    files = []
    for ext in ("*.txt", "*.xlsx", "*.xls", "*.docx", "*.md", "*.csv", "*.json"):
        files.extend(glob.glob(os.path.join(KB_DIR, ext)))
    return sorted(files)


def read_text_file(path: str) -> str:
    """读取纯文本文件"""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试其他编码
        try:
            with open(path, "r", encoding="gbk") as f:
                return f.read()
        except Exception:
            return f"[无法解码文件: {os.path.basename(path)}]"


def read_excel_file(path: str) -> str:
    """读取 Excel 文件"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # 过滤掉全是 None 的行
                vals = [str(v) if v is not None else "" for v in row]
                if any(v.strip() for v in vals):
                    rows.append(" | ".join(vals))
            if rows:
                parts.append(f"【工作表: {sheet_name}】\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(parts) if parts else "(空表格)"
    except Exception as e:
        return f"[读取 Excel 失败: {e}]"


def read_word_file(path: str) -> str:
    """读取 Word 文档"""
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也读取表格
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(c for c in cells):
                    rows.append(" | ".join(cells))
            if rows:
                tables.append("\n".join(rows))
        result = "\n".join(paragraphs)
        if tables:
            result += "\n\n【表格】\n" + "\n\n".join(tables)
        return result if result else "(空文档)"
    except Exception as e:
        return f"[读取 Word 失败: {e}]"


def read_file(path: str) -> str:
    """根据扩展名自动选择读取方式"""
    ext = os.path.splitext(path)[1].lower()
    filename = os.path.basename(path)
    try:
        if ext in (".txt", ".md", ".csv", ".json"):
            content = read_text_file(path)
        elif ext in (".xlsx", ".xls"):
            content = read_excel_file(path)
        elif ext == ".docx":
            content = read_word_file(path)
        else:
            content = "(不支持的格式)"
        
        # 限制单个文件过长
        if len(content) > 10000:
            content = content[:10000] + "\n\n...（内容过长，已截断）"
        return f"📄 {filename}\n{'-'*20}\n{content}"
    except Exception as e:
        return f"[读取 {filename} 失败: {e}]"


@mcp.tool()
def list_knowledge_files() -> str:
    """
    列出知识库中所有可用的文件。
    Bot 可以用这个知道有哪些知识文件可用。
    """
    files = get_supported_files()
    if not files:
        return "知识库为空，请往 knowledge_base 文件夹中添加文件喵~"
    result = ["📚 知识库文件列表："]
    for f in files:
        name = os.path.basename(f)
        size = os.path.getsize(f)
        if size < 1024:
            size_str = f"{size} B"
        elif size < 1024 * 1024:
            size_str = f"{size/1024:.1f} KB"
        else:
            size_str = f"{size/1024/1024:.1f} MB"
        result.append(f"  · {name} ({size_str})")
    return "\n".join(result)


@mcp.tool()
def read_knowledge_file(filename: str) -> str:
    """
    读取知识库中的指定文件内容。
    传入文件名（如 "readme.txt"），Bot 会自动在知识库中查找并读取。
    支持格式：.txt .md .csv .json .xlsx .xls .docx
    """
    files = get_supported_files()
    # 精确匹配
    for f in files:
        if os.path.basename(f).lower() == filename.lower():
            return read_file(f)
    # 模糊匹配
    for f in files:
        if filename.lower() in os.path.basename(f).lower():
            return read_file(f)
    return f"在知识库中找不到「{filename}」喵~ 用 list_knowledge_files 看看有哪些文件？"


@mcp.tool()
def search_knowledge(keyword: str) -> str:
    """
    在知识库所有文件中搜索关键词，返回包含该关键词的片段。
    Bot 可以用这个快速找到相关知识点。
    """
    files = get_supported_files()
    results = []
    keyword_lower = keyword.lower()

    for f in files:
        try:
            content = read_file(f)
            lines = content.split("\n")
            matched_lines = []
            for i, line in enumerate(lines, 1):
                if keyword_lower in line.lower():
                    # 显示匹配行前后各一行
                    start = max(0, i - 2)
                    end = min(len(lines), i + 1)
                    snippet = "\n".join(lines[start:end])
                    matched_lines.append(f"  (行{i}) ...{line.strip()[:150]}...")
                    if len(matched_lines) >= 5:  # 每个文件最多5条匹配
                        break
            if matched_lines:
                name = os.path.basename(f)
                results.append(f"📄 {name}:\n" + "\n".join(matched_lines[:3]))
        except Exception:
            continue

    if not results:
        return f"在知识库中没找到包含「{keyword}」的内容喵~"
    return "🔍 知识库搜索结果：\n\n" + "\n\n".join(results)


@mcp.tool()
def get_knowledge_summary() -> str:
    """
    获取知识库的概览：总文件数、总大小、每个文件的简要说明。
    Bot 可以用这个快速了解知识库全貌。
    """
    files = get_supported_files()
    if not files:
        return "知识库为空，请往 knowledge_base 文件夹中添加文件喵~"

    total_size = sum(os.path.getsize(f) for f in files)
    if total_size < 1024 * 1024:
        size_str = f"{total_size/1024:.1f} KB"
    else:
        size_str = f"{total_size/1024/1024:.1f} MB"

    lines = [f"📚 知识库概览：{len(files)} 个文件，共 {size_str}"]
    for f in files:
        name = os.path.basename(f)
        ext = os.path.splitext(name)[1].lower()
        ext_map = {".txt": "文本", ".md": "Markdown", ".csv": "表格",
                   ".json": "JSON", ".xlsx": "Excel", ".xls": "Excel", ".docx": "Word"}
        file_type = ext_map.get(ext, "文件")
        lines.append(f"  · {name} ({file_type})")
    return "\n".join(lines)


if __name__ == "__main__":
    ensure_kb_dir()
    mcp.run(transport="stdio")
